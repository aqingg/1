import asyncio
import io
import json
import logging
import re
import sys
import zipfile
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import httpx
from fastapi import HTTPException


SERVER_ROOT = Path(__file__).resolve().parents[1]

# Use the server-local vendored python-docx build. It contains the legacy
# datamerge patch that allows macro-enabled .docm templates to be opened.
CUSTOM_DOCX_SRC = SERVER_ROOT / "vendor" / "custom_python_docx" / "src"
if CUSTOM_DOCX_SRC.is_dir():
    custom_docx_src = str(CUSTOM_DOCX_SRC)
    if custom_docx_src not in sys.path:
        sys.path.insert(0, custom_docx_src)
else:
    raise RuntimeError(f"Vendored custom python-docx not found: {CUSTOM_DOCX_SRC}")

import docx  # type: ignore[reportMissingImports]  # noqa: E402

from services.word.logo import replace_logo_placeholders


logger = logging.getLogger("uvicorn.error")


GATEWAY_KEY = "PN9rSrBi6770yG35WSoN25yAPiWaqbBS"
BASE_URL = "http://apiroutecccn.apac.bosch.com/openapi/pmsserverprod/api"

URL_STEP1 = f"{BASE_URL}/getSimpleProjectInfoList?gatewayKey={GATEWAY_KEY}"
URL_STEP2_TEMPLATE = f"{BASE_URL}/vms/get/CustomerProjectByIDX/{{uuid}}?gatewayKey={GATEWAY_KEY}"
URL_STEP3_TEMPLATE = f"{BASE_URL}/vms/pbi/CustomerProjects/ALL?projectId={{uuid}}"


def format_all_roles(team_members: Any) -> str:
    if not isinstance(team_members, list):
        return "N/A"

    role_map: dict[str, list[str]] = defaultdict(list)
    for member in team_members:
        if isinstance(member, dict) and member.get("Role") and member.get("DisplayName"):
            role_map[member["Role"]].append(member["DisplayName"])

    if not role_map:
        return "N/A"

    return "; ".join(
        f"{role}: {', '.join(names)}"
        for role, names in sorted(role_map.items())
    )


def format_all_emails(team_members: Any) -> str:
    if not isinstance(team_members, list):
        return "N/A"

    role_map: dict[str, list[str]] = defaultdict(list)
    for member in team_members:
        if isinstance(member, dict) and member.get("Role") and member.get("DisplayName"):
            usi = member.get("USI")
            if isinstance(usi, dict) and usi.get("Email"):
                member_email = f"{member['DisplayName']}_{usi.get('Email')}"
            else:
                member_email = f"{member['DisplayName']}_NoEmailInfo"
            role_map[member["Role"]].append(member_email)

    if not role_map:
        return "N/A"

    return "; ".join(
        f"{role}Email: {', '.join(names)}"
        for role, names in sorted(role_map.items())
    )


async def fetch_project_identifiers() -> List[Dict[str, str]]:
    async with httpx.AsyncClient() as client:
        try:
            response_step1 = await client.get(URL_STEP1, timeout=30)
            response_step1.raise_for_status()
            projects_data = response_step1.json()
        except (httpx.RequestError, json.JSONDecodeError) as exc:
            raise HTTPException(
                status_code=502,
                detail=f"Failed to fetch PMS project identifiers: {exc}",
            ) from exc

    filtered_projects = []
    if projects_data.get("data") and isinstance(projects_data["data"], list):
        for item in projects_data["data"]:
            if item.get("product_category", "").startswith("AB1"):
                filtered_projects.append({
                    "uuid": item.get("uuid"),
                    "customer": item.get("customer_name"),
                })

    return filtered_projects


async def fetch_single_project_details(uuid: str) -> Optional[Dict[str, Any]]:
    profile: Dict[str, Any] = {
        "customer": "N/A",
        "project": "N/A",
        "ab_generation": "N/A",
        "sop": 0,
        "vehicle_variant": "N/A",
        "plattform": "N/A",
        "type": "N/A",
        "vint_responsible": "N/A",
        "project_leader": "N/A",
        "region": "N/A",
        "oem": "N/A",
        "model": "N/A",
        "peripheral_sensor_configuration": "N/A",
        "internal_sensor_configuration": "N/A",
        "role_summary": "N/A",
        "Digit10OemPn": "N/A",
        "customerOemPn": "N/A",
        "FlConfiguration": "N/A",
        "FlLoopCount": "N/A",
        "TargetMarket": "N/A",
        "MCR_No": "N/A",
        "ConnectorDirection": "N/A",
        "Status": "N/A",
        "AlmProjectName": "N/A",
        "ReferenceProjectPn": "N/A",
        "ReferenceProjectBm": "N/A",
        "role_email_summary": "N/A",
    }

    async def get_step2_data(client: httpx.AsyncClient):
        try:
            response = await client.get(URL_STEP2_TEMPLATE.format(uuid=uuid), timeout=20)
            response.raise_for_status()
            return response.json()
        except (httpx.RequestError, json.JSONDecodeError):
            return None

    async def get_step3_data(client: httpx.AsyncClient):
        try:
            response = await client.get(
                URL_STEP3_TEMPLATE.format(uuid=uuid),
                headers={"gatewayKey": GATEWAY_KEY},
                timeout=300,
            )
            response.raise_for_status()
            return response.json()
        except (httpx.RequestError, json.JSONDecodeError):
            return None

    async with httpx.AsyncClient() as client:
        data2, data3 = await asyncio.gather(
            get_step2_data(client),
            get_step3_data(client),
        )

    if data2:
        profile["customer"] = data2.get("CustomerName", "N/A")
        profile["project"] = (data2.get("VehicleModelNameList") or ["N/A"])[0]
        profile["model"] = (data2.get("VehicleModelNameList") or ["N/A"])[0]
        profile["ab_generation"] = data2.get("ProductCategory", "N/A")
        profile["type"] = data2.get("VehicleSegment", "N/A")
        profile["Digit10OemPn"] = data2.get("Digit10OemPn", "N/A")
        profile["customerOemPn"] = data2.get("customerOemPn", "N/A")
        profile["TargetMarket"] = data2.get("TargetMarket", "N/A")
        profile["Status"] = data2.get("Status", "N/A")
        profile["role_summary"] = format_all_roles(data2.get("TeamMembers", []))

        sop_str = data2.get("TimelineObject", {}).get("CustomerSOP")
        if sop_str:
            try:
                dt_object = datetime.fromisoformat(sop_str.replace("Z", "+00:00"))
                profile["sop"] = int(dt_object.strftime("%Y%m%d"))
            except (ValueError, TypeError):
                profile["sop"] = 0

        profile["plattform"] = (data2.get("PlatformList") or ["N/A"])[0]
        profile["region"] = data2.get("respRegion", "N/A")
        profile["vehicle_variant"] = (data2.get("VehicleModelNameList") or ["N/A"])[0]

    if data3:
        profile["oem"] = data3.get("CustomerName", "N/A")
        profile["vint_responsible"] = data3.get("RespSWPCM", "N/A")
        profile["project_leader"] = data3.get("RespTPM", "N/A")
        profile["FlConfiguration"] = data3.get("FlConfiguration", "N/A")
        profile["internal_sensor_configuration"] = data3.get("InternalSensor", "N/A")
        profile["MCR_No"] = data3.get("MCR_L0", "N/A")
        profile["ConnectorDirection"] = data3.get("ConnectorDirection", "N/A")
        profile["AlmProjectName"] = data3.get("AlmProjectName", "N/A")
        profile["ReferenceProjectPn"] = data3.get("ReferenceProjectPn", "N/A")
        profile["ReferenceProjectBm"] = data3.get("ReferenceProjectBm", "N/A")
        profile["role_email_summary"] = format_all_emails(data3.get("TeamMembers", []))

        all_sensors = [data3.get(key, "") for key in ["Ufs", "Pas", "Pps"]]
        valid_parts = [sensor for sensor in all_sensors if sensor and sensor != "0"]
        profile["peripheral_sensor_configuration"] = (
            "+".join(valid_parts) if valid_parts else "N/A"
        )

    for key, value in profile.items():
        if value == "null" or value is None:
            profile[key] = "N/A"

    profile["FlLoopCount"] = format_firing_loop_count(profile.get("FlConfiguration"))

    return profile


def _parse_role_email_summary(role_email_summary: str) -> dict[str, str]:
    if not role_email_summary or role_email_summary == "N/A":
        return {}

    email_data_by_role = {}
    for section in role_email_summary.split(";"):
        if ":" not in section:
            continue
        role_key, _, member_list_str = section.partition(":")
        role_key = role_key.strip()
        if not role_key:
            continue

        output_lines = []
        for member_str in member_list_str.strip().split(","):
            member_str = member_str.strip()
            if not member_str:
                continue
            display_name, separator, email = member_str.rpartition("_")
            if separator:
                output_lines.append(display_name.strip())
                output_lines.append(email.strip())

        if output_lines:
            email_data_by_role[role_key] = "\n".join(output_lines)

    return email_data_by_role


def format_firing_loop_count(value: Any) -> str:
    if isinstance(value, (list, tuple, set)):
        raw_value = "+".join(str(item).strip() for item in value if str(item).strip())
    else:
        raw_value = str(value or "").strip()

    if not raw_value or raw_value.upper() in {"N/A", "NA", "NULL", "NONE", "0"}:
        return "N/A"

    existing_count = re.fullmatch(r"(\d+)\s*loops?", raw_value, flags=re.IGNORECASE)
    if existing_count:
        return existing_count.group(1)

    loops = [
        part.strip()
        for part in re.split(r"[+,\n;]+", raw_value)
        if part.strip()
    ]
    if not loops:
        return "N/A"

    return str(len(loops))


def _prepare_profile_for_filling(profile_dict: dict[str, Any]) -> dict[str, str]:
    formatted_profile = {}
    for key, value in profile_dict.items():
        if value is None or value == "" or value == []:
            formatted_profile[key] = "N/A"
        elif isinstance(value, list):
            formatted_profile[key] = ", ".join(map(str, value))
        else:
            formatted_profile[key] = str(value)

    formatted_profile["FlLoopCount"] = format_firing_loop_count(
        formatted_profile.get("FlConfiguration")
    )

    role_summary_str = formatted_profile.get("role_summary", "N/A")
    role_data = {}
    if role_summary_str and role_summary_str != "N/A":
        for pair in role_summary_str.split(";"):
            if ":" in pair:
                role, _, names = pair.partition(":")
                role_data[role.strip()] = names.strip()

    role_mapping = {
        "PJM": "PjM",
        "TPM": "TPM",
        "ECU_PCM": "ECU-PCM",
        "SW_PCM": "SW_PCM",
        "FSM": "FSM",
        "SYS_ENG": "Sys-ENG",
        "APP_PCM": "App PCM",
        "HW_Dev": "HW Developer",
        "AM": "AM",
        "CM": "CM",
        "COS": "COS",
        "MECH_PCM": "MECH-PCM",
        "SAMCO": "SAMCO",
        "SEC": "SEC",
        "TestM": "Test Manager",
    }

    for placeholder_key, summary_key in role_mapping.items():
        formatted_profile[placeholder_key] = role_data.get(summary_key, "N/A")

    inverted_role_mapping = {value: key for key, value in role_mapping.items()}
    parsed_emails_by_role = _parse_role_email_summary(
        formatted_profile.get("role_email_summary", "N/A")
    )
    corrected_email_data = {}
    for raw_key, value in parsed_emails_by_role.items():
        if raw_key.endswith("Email"):
            prefix = raw_key[:-5]
            placeholder_prefix = inverted_role_mapping.get(prefix)
            corrected_email_data[
                f"{placeholder_prefix}Email" if placeholder_prefix else raw_key
            ] = value
        else:
            corrected_email_data[raw_key] = value

    formatted_profile.update(corrected_email_data)
    return formatted_profile


def fill_docx_by_placeholders(
    profile_dict: dict[str, Any],
    source: Union[Path, io.BytesIO],
) -> io.BytesIO:
    try:
        document = docx.Document(source)
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"无法处理提供的Word文件: {exc}") from exc

    formatted_profile = _prepare_profile_for_filling(profile_dict)
    placeholder_pattern = re.compile(r"<\s*PMS\.([^>]+?)\s*>")

    def normalize_header_text(value: str) -> str:
        return re.sub(r"\s+", " ", (value or "").strip()).lower()

    def header_index_map(header_row) -> Dict[str, int]:
        mapping: Dict[str, int] = {}
        for index, cell in enumerate(header_row.cells):
            key = normalize_header_text(cell.text)
            if key and key not in mapping:
                mapping[key] = index
        return mapping

    def set_cell_text(row, column_index: int, text: str) -> None:
        if column_index >= len(row.cells):
            return
        row.cells[column_index].text = text

    def normalize_document_change_table(document) -> None:
        required_headers = [
            "no",
            "document",
            "author",
            "description of change",
            "version",
            "date",
        ]

        for table in document.tables:
            if not table.rows:
                continue

            header_map = header_index_map(table.rows[0])
            if not all(header in header_map for header in required_headers):
                continue

            # 保留模板里已经写好的第一条变更记录，只清理多余行。
            while len(table.rows) > 2:
                table._tbl.remove(table.rows[-1]._tr)
            break

    normalize_document_change_table(document)

    def replacer(match: re.Match) -> str:
        combined_key = match.group(1).strip()
        if combined_key.lower() == "logo":
            return "<PMS.logo>"
        if "-" in combined_key:
            keys = [key.strip() for key in combined_key.split("-")]
            value_parts = [formatted_profile.get(key, "N/A") for key in keys]
            return "_".join(value_parts)
        return formatted_profile.get(combined_key, "N/A")

    def substitute_in_paragraph(paragraph):
        if placeholder_pattern.search(paragraph.text):
            full_text = paragraph.text
            new_text = placeholder_pattern.sub(replacer, full_text)
            if new_text != full_text:
                paragraph.text = new_text

    def substitute_in_container(container):
        for paragraph in container.paragraphs:
            substitute_in_paragraph(paragraph)

        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    substitute_in_container(cell)

    substitute_in_container(document)

    for section in document.sections:
        substitute_in_container(section.header)
        substitute_in_container(section.first_page_header)
        substitute_in_container(section.even_page_header)
        substitute_in_container(section.footer)
        substitute_in_container(section.first_page_footer)
        substitute_in_container(section.even_page_footer)

    logo_replacements = replace_logo_placeholders(document, formatted_profile)
    if logo_replacements:
        logger.info("[TCD08] Replaced %s logo placeholder(s).", logo_replacements)

    stream = io.BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def flatten_project_info(project_info: Dict[str, Any]) -> Dict[str, str]:
    values = {}

    def add_value(label, value):
        if label and value not in [None, "", []]:
            values[label] = ", ".join(map(str, value)) if isinstance(value, list) else str(value)

    for key in ["owner", "proxies", "uuid"]:
        item = project_info.get(key)
        if isinstance(item, dict):
            add_value(item.get("label"), item.get("value"))

    for row in project_info.get("projectInfo", []) or []:
        if not isinstance(row, list):
            continue
        for item in row:
            if isinstance(item, dict):
                add_value(item.get("label"), item.get("value"))

    return values


def apply_project_info_overrides(
    profile_dict: Dict[str, Any],
    project_info: Dict[str, Any],
) -> Dict[str, Any]:
    form_values = flatten_project_info(project_info)
    form_to_profile = {
        "OEM": "oem",
        "Product Category": "ab_generation",
        "Market": "TargetMarket",
        "Status": "Status",
        "SOP Date": "sop",
        "Project Leader": "project_leader",
        "MCR No.": "MCR_No",
        "ECU Direction": "ConnectorDirection",
        "BOSCH PIN": "Digit10OemPn",
        "Customer PIN": "customerOemPn",
        "Inertial Sensor": "internal_sensor_configuration",
        "Peripheral Sensor": "peripheral_sensor_configuration",
        "Vehicle Type": "type",
        "Fire Loops": "FlConfiguration",
    }

    for form_label, profile_key in form_to_profile.items():
        if form_values.get(form_label):
            profile_dict[profile_key] = form_values[form_label]

    if form_values.get("Owner"):
        profile_dict["author"] = form_values["Owner"]
    elif form_values.get("Project Leader"):
        profile_dict["author"] = form_values["Project Leader"]

    profile_dict["FlLoopCount"] = format_firing_loop_count(
        profile_dict.get("FlConfiguration")
    )

    return profile_dict

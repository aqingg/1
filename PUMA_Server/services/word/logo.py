from __future__ import annotations

import logging
from pathlib import Path
from typing import Any

from docx.image.exceptions import UnrecognizedImageError
from docx.image.image import Image
from docx.shared import Inches

from services.word.xml_utils import clean_text
from utils.file_loader import load_data_source
from utils.path_config import BASE_RUNTIME_DIR, CUSTOMER_LOGO_DIR, DATA_SOURCE_DIR, PROJECT_ROOT


logger = logging.getLogger("uvicorn.error")

LOGO_PLACEHOLDER = "<PMS.logo>"
LOGO_MAPPING_FILENAME = "LogoMapping.json"


def _normalized_key(value: Any) -> str:
    return clean_text(value).strip().lower()


def _to_inch_value(value: Any):
    try:
        if value is None or value == "":
            return None
        return Inches(float(value))
    except (TypeError, ValueError):
        return None


def _resolve_image_path(path_value: str) -> Path | None:
    candidate = Path(path_value)
    if candidate.is_absolute():
        return candidate if candidate.exists() else None

    search_roots = [CUSTOMER_LOGO_DIR, BASE_RUNTIME_DIR, DATA_SOURCE_DIR, PROJECT_ROOT]
    for root in search_roots:
        resolved = (root / candidate).resolve()
        if resolved.exists():
            return resolved

    return None


def _is_supported_image(image_path: Path) -> bool:
    try:
        Image.from_file(str(image_path))
        return True
    except (FileNotFoundError, UnrecognizedImageError, OSError):
        return False


def _load_logo_mapping() -> dict[str, Any]:
    try:
        mapping = load_data_source(LOGO_MAPPING_FILENAME)
    except FileNotFoundError:
        logger.info("[TCD08] Logo mapping file not found: %s", LOGO_MAPPING_FILENAME)
        return {}
    except Exception as exc:
        logger.warning("[TCD08] Failed to load logo mapping: %s", exc, exc_info=True)
        return {}

    return mapping if isinstance(mapping, dict) else {}


def _logo_candidate_keys(profile_dict: dict[str, Any]) -> list[str]:
    candidate_keys = []
    for value in (
        profile_dict.get("oem"),
        profile_dict.get("customer"),
        profile_dict.get("project"),
        profile_dict.get("customer_name"),
        profile_dict.get("customerName"),
    ):
        normalized = _normalized_key(value)
        if normalized and normalized not in candidate_keys:
            candidate_keys.append(normalized)

    if "default" not in candidate_keys:
        candidate_keys.append("default")

    return candidate_keys


def _coerce_logo_spec(value: Any) -> dict[str, Any] | None:
    if isinstance(value, str):
        return {"path": value}
    if isinstance(value, dict) and value.get("path"):
        return value
    return None


def _candidate_matches_spec(candidate_key: str, spec_key: str, spec: dict[str, Any]) -> bool:
    if not candidate_key or not spec_key:
        return False

    normalized_spec_key = _normalized_key(spec_key)
    if candidate_key == normalized_spec_key:
        return True

    if candidate_key in normalized_spec_key or normalized_spec_key in candidate_key:
        return True

    aliases = spec.get("aliases")
    if isinstance(aliases, (list, tuple, set)):
        for alias in aliases:
            normalized_alias = _normalized_key(alias)
            if not normalized_alias:
                continue
            if candidate_key == normalized_alias:
                return True
            if candidate_key in normalized_alias or normalized_alias in candidate_key:
                return True

    return False


def resolve_logo_spec(profile_dict: dict[str, Any]) -> dict[str, Any] | None:
    mapping = _load_logo_mapping()
    for candidate_key in _logo_candidate_keys(profile_dict):
        normalized_candidate = _normalized_key(candidate_key)
        for spec_key, raw_spec in mapping.items():
            if _normalized_key(spec_key) == "default":
                continue
            spec = _coerce_logo_spec(raw_spec)
            if spec is None:
                continue
            if _candidate_matches_spec(normalized_candidate, spec_key, spec):
                return spec

    default_spec = _coerce_logo_spec(mapping.get("default"))
    if default_spec is not None:
        return default_spec
    return None


def _is_logo_placeholder_paragraph(paragraph) -> bool:
    return _normalized_key(getattr(paragraph, "text", "")) == _normalized_key(LOGO_PLACEHOLDER)


def _replace_logo_placeholder_in_container(container, image_path: Path, width=None, height=None) -> int:
    replacements = 0

    for paragraph in container.paragraphs:
        if _is_logo_placeholder_paragraph(paragraph):
            paragraph.clear()
            run = paragraph.add_run()
            if width is None and height is None:
                run.add_picture(str(image_path))
            else:
                run.add_picture(str(image_path), width=width, height=height)
            replacements += 1

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                replacements += _replace_logo_placeholder_in_container(cell, image_path, width, height)

    return replacements


def replace_logo_placeholders(document, profile_dict: dict[str, Any]) -> int:
    """Replace standalone <PMS.logo> paragraphs with a DB-managed image.

    The logo placeholder is expected to be a standalone paragraph or cell content.
    The image source is resolved from LogoMapping.json, which points to files under
    BASE_RUNTIME_DIR/Customer_Logo or other project-relative locations.
    """
    spec = resolve_logo_spec(profile_dict)
    if spec is None:
        return 0

    image_path_value = str(spec.get("path") or "").strip()
    if not image_path_value:
        logger.warning("[TCD08] Logo mapping resolved without a path: %s", spec)
        return 0

    image_path = _resolve_image_path(image_path_value)
    if image_path is None:
        logger.warning("[TCD08] Logo image not found: %s", image_path_value)
        return 0

    if not _is_supported_image(image_path):
        logger.warning("[TCD08] Logo image format is not supported by python-docx: %s", image_path)
        return 0

    width = _to_inch_value(spec.get("width_inches"))
    height = _to_inch_value(spec.get("height_inches"))

    replacements = 0
    replacements += _replace_logo_placeholder_in_container(document, image_path, width, height)

    for section in document.sections:
        replacements += _replace_logo_placeholder_in_container(section.header, image_path, width, height)
        replacements += _replace_logo_placeholder_in_container(section.first_page_header, image_path, width, height)
        replacements += _replace_logo_placeholder_in_container(section.even_page_header, image_path, width, height)
        replacements += _replace_logo_placeholder_in_container(section.footer, image_path, width, height)
        replacements += _replace_logo_placeholder_in_container(section.first_page_footer, image_path, width, height)
        replacements += _replace_logo_placeholder_in_container(section.even_page_footer, image_path, width, height)

    return replacements